"""Extracts text from uploaded legal documents (PDF/DOCX) and asks Claude to
review them: clause extraction, risk assessment, obligations, missing terms,
and a confidence-scored summary."""
import os

from pypdf import PdfReader
from docx import Document as DocxDocument
from anthropic import Anthropic

# Sonnet 5 for stronger multi-step legal reasoning than a fast/cheap model
# would give — a single analysis costs a few cents to ~tens of cents
# depending on document length. Swap to claude-haiku-4-5 to cut cost.
MODEL = "claude-sonnet-5"
# Sonnet 5 has a 200K-token context window shared by input + output. These
# caps sit near the top of what fits alongside the schema/prompt overhead and
# the reserved 8192-token output — uploads can total up to 800MB on disk (see
# MAX_CONTENT_LENGTH in app.py), but only the text below actually reaches the
# model; anything past MAX_TOTAL_CHARS is dropped and flagged via `truncated`
# in the extract_document_text result, since most of an 800MB PDF set is
# images/fonts, not extractable text, and text volume is what the context
# window actually limits.
MAX_CHARS_PER_DOC = 350_000
MAX_TOTAL_CHARS = 600_000

CLAUSE_CATEGORIES = [
    "Indemnification", "Limitation of Liability", "Termination",
    "Assignment & Change of Control", "Governing Law & Venue",
    "Confidentiality", "Non-Compete / Non-Solicit",
    "Intellectual Property Ownership", "Payment Terms",
    "Renewal / Auto-Renewal", "Force Majeure",
    "Dispute Resolution / Arbitration", "Insurance Requirements",
    "Representations & Warranties", "Default & Remedies", "Exclusivity",
]

ANALYZE_SCHEMA = {
    "name": "analyze_legal_document",
    "description": "Report a structured legal review of the provided document(s).",
    "input_schema": {
        "type": "object",
        "properties": {
            "document_type": {
                "type": "string",
                "description": "Best guess at the document type, e.g. 'Commercial Lease Agreement', 'Purchase and Sale Agreement', 'Credit/Loan Agreement', 'NDA'.",
            },
            "parties": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string"},
                        "role": {"type": "string", "description": "e.g. 'Landlord', 'Tenant', 'Borrower', 'Lender', 'Seller', 'Buyer'"},
                    },
                    "required": ["name", "role"],
                },
            },
            "key_dates": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "label": {"type": "string", "description": "e.g. 'Effective Date', 'Expiration Date', 'Renewal Notice Deadline'"},
                        "value": {"type": "string"},
                    },
                    "required": ["label", "value"],
                },
            },
            "summary": {
                "type": "string",
                "description": "3-5 sentence plain-English overview of what the document does and its overall risk posture.",
            },
            "overall_confidence": {
                "type": "integer",
                "description": "0-100 confidence in the completeness/accuracy of this analysis, given the document's clarity and completeness.",
            },
            "clauses": {
                "type": "array",
                "description": "One entry per standard clause category that is actually present in the document. Skip categories that don't apply.",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {"type": "string", "description": f"One of: {', '.join(CLAUSE_CATEGORIES)} (or a close custom label if none fit)"},
                        "location_hint": {"type": "string", "description": "Section/article name or number if identifiable, e.g. 'Section 8.2'"},
                        "explanation": {"type": "string", "description": "Plain-English explanation of what this clause says and what it means for the reviewing party."},
                        "risk_level": {"type": "string", "enum": ["high", "medium", "low", "none"]},
                        "risk_note": {"type": "string", "description": "Why this risk level was assigned, or empty string if risk_level is 'none'."},
                        "confidence": {"type": "integer", "description": "0-100 confidence in this specific clause read."},
                    },
                    "required": ["category", "explanation", "risk_level", "confidence"],
                },
            },
            "obligations": {
                "type": "array",
                "description": "Concrete obligations/covenants each party owes under the document.",
                "items": {
                    "type": "object",
                    "properties": {
                        "party": {"type": "string"},
                        "obligation": {"type": "string"},
                        "trigger_or_deadline": {"type": "string", "description": "When/what triggers it, or empty string if ongoing/unconditional."},
                    },
                    "required": ["party", "obligation"],
                },
            },
            "missing_terms": {
                "type": "array",
                "description": "Standard clauses/protections a document of this type would typically include but that appear absent here.",
                "items": {
                    "type": "object",
                    "properties": {
                        "term": {"type": "string"},
                        "why_it_matters": {"type": "string"},
                        "recommendation": {"type": "string"},
                    },
                    "required": ["term", "why_it_matters", "recommendation"],
                },
            },
        },
        "required": ["document_type", "summary", "overall_confidence", "clauses", "obligations", "missing_terms"],
    },
}


class AnalyzerError(Exception):
    pass


def _extract_pdf(fs):
    try:
        reader = PdfReader(fs.stream)
    except Exception as e:
        raise AnalyzerError(f"Could not read '{fs.filename}': it may not be a valid PDF.") from e
    pages_text = []
    total_len = 0
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            continue
        pages_text.append(text)
        total_len += len(text)
        if total_len > MAX_CHARS_PER_DOC:
            break
    text = "\n".join(pages_text).strip()
    if not text:
        raise AnalyzerError(
            f"Couldn't extract any text from '{fs.filename}' — it may be a scanned image "
            "PDF with no text layer, which this tool can't read (no OCR)."
        )
    return text


def _extract_docx(fs):
    try:
        doc = DocxDocument(fs.stream)
    except Exception as e:
        raise AnalyzerError(f"Could not read '{fs.filename}': it may not be a valid .docx file.") from e
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_document_text(file_storage_list):
    """file_storage_list: list of werkzeug FileStorage objects (.pdf or .docx).
    Returns (text, truncated) — truncated is True if the combined text had to
    be cut short to fit the model's context window."""
    chunks = []
    total = 0
    truncated = False
    for fs in file_storage_list:
        if not fs or not fs.filename:
            continue
        name_lower = fs.filename.lower()
        if name_lower.endswith(".pdf"):
            doc_text = _extract_pdf(fs)
        elif name_lower.endswith(".docx"):
            doc_text = _extract_docx(fs)
        elif name_lower.endswith(".doc"):
            raise AnalyzerError(
                f"'{fs.filename}' is an old-format .doc file — please save it as .docx or PDF and re-upload."
            )
        else:
            raise AnalyzerError(f"Unsupported file type: '{fs.filename}'. Only .pdf and .docx files are supported.")

        if len(doc_text) > MAX_CHARS_PER_DOC:
            doc_text = doc_text[:MAX_CHARS_PER_DOC]
            truncated = True
        if not doc_text:
            continue
        chunk = f"\n\n===== DOCUMENT: {fs.filename} =====\n{doc_text}"
        if total + len(chunk) > MAX_TOTAL_CHARS:
            chunk = chunk[: max(0, MAX_TOTAL_CHARS - total)]
            truncated = True
        chunks.append(chunk)
        total += len(chunk)
        if total >= MAX_TOTAL_CHARS:
            break
    return "".join(chunks).strip(), truncated


def analyze_document(document_text: str) -> dict:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise AnalyzerError(
            "No Anthropic API key configured. Add ANTHROPIC_API_KEY to the .env file "
            "in the legal-doc-analyzer folder and restart the server."
        )
    if not document_text:
        raise AnalyzerError("Couldn't extract any readable text from the uploaded file(s).")

    client = Anthropic(api_key=api_key)

    prompt = f"""You are a commercial contract reviewer preparing an internal risk-review memo. \
Read the attached document(s) and call the analyze_legal_document tool with a thorough, \
well-supported review.

Rules:
- Base every field only on what the document(s) actually say. Do not invent parties, dates, \
figures, or clause language that isn't there.
- Write in plain English a business reader (not a lawyer) can act on — explain what each clause \
means practically, not just what it says.
- clauses: cover every standard category from this list that is actually present in the document \
— {', '.join(CLAUSE_CATEGORIES)} — plus any other clause worth flagging. Skip categories that \
genuinely don't apply; don't force an entry for every category.
- risk_level reflects how unfavorable/unusual the term is for the party reviewing this document \
(assume the reviewer is whichever party seems to be receiving/signing rather than drafting, if \
that's discernible; otherwise assess neutrally). "none" means the clause is standard/market and \
poses no notable risk.
- missing_terms: flag standard protections typical for this document type that appear absent — \
e.g. no limitation of liability cap, no indemnification carve-outs, no assignment restriction, \
no insurance requirement. Only list terms that would genuinely be expected for this document type.
- confidence scores should reflect real uncertainty — lower them when the document is ambiguous, \
uses unusual defined terms, or is missing context (e.g. referenced exhibits not included).

DOCUMENT(S):
{document_text}
"""

    try:
        resp = client.messages.create(
            model=MODEL,
            max_tokens=8192,
            tools=[ANALYZE_SCHEMA],
            tool_choice={"type": "tool", "name": "analyze_legal_document"},
            messages=[{"role": "user", "content": prompt}],
        )
    except Exception as e:
        raise AnalyzerError(f"Claude API request failed: {e}") from e

    for block in resp.content:
        if block.type == "tool_use" and block.name == "analyze_legal_document":
            return block.input

    raise AnalyzerError("Claude did not return a structured analysis. Try again.")


CHAT_INSTRUCTIONS = """You are a helpful assistant answering follow-up questions about the \
document(s) below, for someone who has already seen an automated clause/risk review of them.

Rules:
- Answer only from what the document(s) actually say. If the answer isn't in the text, say so \
plainly rather than guessing or inventing specifics.
- Reference section/article numbers or defined terms when it helps the reader locate the answer.
- Keep answers concise and direct — a few sentences unless the question genuinely needs more.
- This is not legal advice; note that if the question is asking you to make a legal judgment call \
rather than report what the document says."""


def chat_about_document(document_text: str, question: str, history: list) -> str:
    """history: list of {"role": "user"|"assistant", "content": str} prior turns (not including
    the new question). The document text is sent as a cached system block so repeated follow-up
    questions on the same document don't reprocess it at full cost each time."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise AnalyzerError(
            "No Anthropic API key configured. Add ANTHROPIC_API_KEY to the .env file "
            "in the legal-doc-analyzer folder and restart the server."
        )
    if not document_text:
        raise AnalyzerError("No document text available to answer questions about — try re-analyzing the document.")
    if not question or not question.strip():
        raise AnalyzerError("Please enter a question.")

    client = Anthropic(api_key=api_key)

    system = [
        {"type": "text", "text": CHAT_INSTRUCTIONS},
        {
            "type": "text",
            "text": f"DOCUMENT(S):\n{document_text}",
            "cache_control": {"type": "ephemeral"},
        },
    ]

    messages = []
    for turn in history or []:
        role = turn.get("role")
        content = turn.get("content")
        if role in ("user", "assistant") and content:
            messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": question})

    try:
        resp = client.messages.create(
            model=MODEL,
            max_tokens=1536,
            system=system,
            messages=messages,
        )
    except Exception as e:
        raise AnalyzerError(f"Claude API request failed: {e}") from e

    text_parts = [b.text for b in resp.content if b.type == "text"]
    answer = "".join(text_parts).strip()
    if not answer:
        raise AnalyzerError("Claude did not return an answer. Try again.")
    return answer
